from docx import Document
from deep_translator import GoogleTranslator
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

input_file = "SAP建Forwarder與Express--程式BP_20250528.docx"
output_file = "SAP建Forwarder與Express--程式BP_20250528_含翻譯.docx"

def split_sentences(text):
    return [s.strip() for s in re.split(r'(?<=[。！？!?\.])', text) if s.strip()]

def translate_text(text):
    try:
        return GoogleTranslator(source='zh-TW', target='en').translate(text)
    except Exception as e:
        print(f"⚠️ 翻譯錯誤: {e}")
        return "[Translation failed]"

def insert_paragraph_after(paragraph: Paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if style:
        new_para.style = style
    run = new_para.add_run(text)
    return new_para

def translate_docx_inplace(input_path, output_path):
    doc = Document(input_path)

    # 建立原始段落副本
    original_paragraphs = list(doc.paragraphs)

    for para in original_paragraphs:
        if not para.text.strip():
            continue  # 忽略空段落

        sentences = split_sentences(para.text)
        for sentence in sentences:
            translated = translate_text(sentence)
            insert_paragraph_after(para, f"{{EN: {translated}}}", style=para.style)

    doc.save(output_path)
    print(f"\n✅ 完成！已儲存翻譯檔案：{output_path}")

# 執行翻譯
translate_docx_inplace(input_file, output_file)
